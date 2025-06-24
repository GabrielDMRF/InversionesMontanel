from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
import os
import re
from datetime import datetime

app = Flask(__name__)
app.secret_key = 'tu_clave_secreta_aqui'

# Configuración de carpetas
TEMPLATE_FOLDER = 'templates_word'
OUTPUT_FOLDER = 'output'

# Crear carpetas si no existen
os.makedirs(TEMPLATE_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Nombre del archivo plantilla (debe estar en la carpeta templates_word)
TEMPLATE_FILE = 'ACUERDO_DE_SEGURIDAD_PLANTILLA.docx'

def procesar_documento_word(datos_formulario):
    """
    Procesa el documento plantilla reemplazando las variables con los datos del formulario
    """
    try:
        # Ruta del archivo plantilla
        template_path = os.path.join(TEMPLATE_FOLDER, TEMPLATE_FILE)
        
        # Verificar que existe la plantilla
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_FILE}")
        
        # Abrir el documento plantilla
        doc = Document(template_path)
        
        # Obtener fecha actual
        fecha_actual = datetime.now()
        dia = fecha_actual.day
        meses = [
            'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
            'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
        ]
        mes = meses[fecha_actual.month - 1]
        año = fecha_actual.year
        
        # Diccionario de reemplazos
        reemplazos = {
            '{{razon_social}}': datos_formulario.get('razon_social', ''),
            '{{nit}}': datos_formulario.get('nit', ''),
            '{{representante}}': datos_formulario.get('representante', ''),
            '{{direccion}}': datos_formulario.get('direccion', ''),
            '{{telefono}}': datos_formulario.get('telefono', ''),
            '{{cedula}}': datos_formulario.get('cedula', ''),
            # Reemplazos para la fecha
            '{{dia}}': str(dia),
            '{{mes}}': mes,
            '{{año}}': str(año)
        }
        
        # Función para reemplazar la fecha usando regex (si usas el formato original)
        def reemplazar_fecha_guiones(texto):
            # Patrón: "fecha del ___ de ________ del 20_____"
            patron_fecha = r'fecha del\s+_{2,}\s+de\s+_{4,}\s+del\s+20_{2,}'
            fecha_completa = f"fecha del {dia} de {mes} del {año}"
            return re.sub(patron_fecha, fecha_completa, texto)
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            # Reemplazar plantillas con {{}}
            for placeholder, valor in reemplazos.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, valor)
            
            # Reemplazar fecha con guiones (formato original)
            paragraph.text = reemplazar_fecha_guiones(paragraph.text)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Reemplazos básicos
                    for placeholder, valor in reemplazos.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, valor)
                    
                    # Reemplazar fecha con guiones
                    cell.text = reemplazar_fecha_guiones(cell.text)
        
        # Generar nombre de archivo único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        razon_social_clean = datos_formulario.get('razon_social', 'empresa').replace(' ', '_').replace('.', '')
        nombre_salida = f"Acuerdo_Seguridad_{razon_social_clean}_{timestamp}.docx"
        ruta_salida = os.path.join(OUTPUT_FOLDER, nombre_salida)
        
        # Guardar el documento modificado
        doc.save(ruta_salida)
        
        return ruta_salida, nombre_salida
        
    except Exception as e:
        print(f"Error procesando documento: {str(e)}")
        return None, None

@app.route('/')
def index():
    # Verificar si existe la plantilla
    template_path = os.path.join(TEMPLATE_FOLDER, TEMPLATE_FILE)
    template_exists = os.path.exists(template_path)
    
    return render_template('formulario.html', template_exists=template_exists)

@app.route('/generar', methods=['POST'])
def generar_documento():
    try:
        # Obtener datos del formulario
        datos_formulario = {
            'razon_social': request.form.get('razon_social', '').strip(),
            'nit': request.form.get('nit', '').strip(),
            'representante': request.form.get('representante', '').strip(),
            'direccion': request.form.get('direccion', '').strip(),
            'telefono': request.form.get('telefono', '').strip(),
            'cedula': request.form.get('cedula', '').strip()
        }
        
        # Validar campos obligatorios
        campos_obligatorios = ['razon_social', 'nit', 'representante', 'direccion', 'telefono', 'cedula']
        campos_vacios = [campo for campo in campos_obligatorios if not datos_formulario[campo]]
        
        if campos_vacios:
            flash(f'Los siguientes campos son obligatorios: {", ".join(campos_vacios)}', 'error')
            return redirect(url_for('index'))
        
        # Procesar el documento
        ruta_salida, nombre_archivo = procesar_documento_word(datos_formulario)
        
        if ruta_salida:
            fecha_actual = datetime.now().strftime("%d de %B del %Y")
            flash(f'Documento generado exitosamente con fecha: {fecha_actual}', 'success')
            return send_file(ruta_salida, as_attachment=True, download_name=nombre_archivo)
        else:
            flash('Error generando el documento. Verifique que la plantilla existe.', 'error')
            return redirect(url_for('index'))
            
    except Exception as e:
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/admin')
def admin():
    """Página de administración para verificar archivos"""
    template_path = os.path.join(TEMPLATE_FOLDER, TEMPLATE_FILE)
    template_exists = os.path.exists(template_path)
    
    output_files = []
    if os.path.exists(OUTPUT_FOLDER):
        output_files = [f for f in os.listdir(OUTPUT_FOLDER) if f.endswith('.docx')]
    
    return render_template('admin.html', 
                         template_exists=template_exists,
                         template_path=template_path,
                         output_files=output_files)

if __name__ == '__main__':
    app.run(debug=True)